# engine.py
"""
DocAI Trainer - Engine
Core processing logic: file extraction, chunking, merging, session management.
Improvements:
  - Support PDF, DOC, DOCX, JPG, PNG, JPEG, WEBP, GIF images (OCR)
  - Up to 200 files, 200 MB per file
  - Parallel processing with ThreadPoolExecutor
  - Auto-paraphrase on process click (without API)
  - Default paraphrase language = auto-detect
"""

import os
import json
import uuid
import shutil
import time
import re
import random
import urllib.request
import urllib.error
from concurrent.futures import ThreadPoolExecutor, as_completed
from datetime import datetime, timezone
from pathlib import Path
from typing import Optional

# ──────────────────────────────────────────────────────────────────────────────
# CONSTANTS
# ──────────────────────────────────────────────────────────────────────────────

UPLOADS_DIR = Path("uploads")
UPLOADS_DIR.mkdir(exist_ok=True)

ALLOWED_EXTENSIONS = {".pdf", ".doc", ".docx", ".txt", ".md", ".jpg", ".jpeg", ".png", ".webp", ".gif"}
IMAGE_EXTENSIONS = {".jpg", ".jpeg", ".png", ".webp", ".gif"}
MAX_FILE_SIZE_MB = 200
MAX_FILES = 200

# Indonesian synonyms for paraphrasing
INDONESIAN_SYNONYMS = {
    "menggunakan": ["memakai", "mengaplikasikan", "menerapkan"],
    "membuat": ["menyusun", "merancang", "menghasilkan"],
    "mendapatkan": ["memperoleh", "meraih", "menerima"],
    "melakukan": ["menjalankan", "mengerjakan", "mengaplikasikan"],
    "mempunyai": ["memiliki", "memperoleh", "mengandung"],
    "menunjukkan": ["menampilkan", "memperlihatkan", "mengindikasikan"],
    "menghasilkan": ["memproduksi", "menciptakan", "menyediakan"],
    "mengurangi": ["meminimalkan", "menurunkan", "memperkecil"],
    "meningkatkan": ["menambah", "memperbesar", "mengembangkan"],
    "mengembangkan": ["memajukan", "memperluas", "meningkatkan"],
    "penting": ["signifikan", "krusial", "esensial"],
    "baik": ["bagus", "berkualitas", "memadai"],
    "besar": ["luas", "signifikan", "substansial"],
    "kecil": ["minimal", "sedikit", "terbatas"],
    "cepat": ["sigap", "efisien", "responsif"],
    "mudah": ["sederhana", "praktis", "tidak rumit"],
    "sulit": ["kompleks", "menantang", "rumit"],
    "juga": ["pun", "lagipula", "selain itu"],
    "karena": ["sebab", "akibat", "lantaran"],
    "tetapi": ["namun", "akan tetapi", "meskipun demikian"],
    "dengan": ["melalui", "lewat", "menggunakan"],
    "untuk": ["bagi", "kepada", "demi"],
    "dari": ["asal", "berasal", "pangkal"],
    "dalam": ["di", "pada", "ke dalam"],
    "atau": ["maupun", "ataupun", "bahkan"],
    "dan": ["serta", "dan juga", "lagi pula"],
}

ENGLISH_SYNONYMS = {
    "use": ["utilize", "employ", "apply"],
    "make": ["create", "construct", "build"],
    "get": ["obtain", "acquire", "receive"],
    "do": ["perform", "execute", "conduct"],
    "have": ["possess", "contain", "hold"],
    "show": ["display", "demonstrate", "indicate"],
    "produce": ["generate", "create", "yield"],
    "reduce": ["minimize", "decrease", "lessen"],
    "increase": ["enhance", "boost", "improve"],
    "develop": ["advance", "expand", "grow"],
    "important": ["significant", "crucial", "essential"],
    "good": ["excellent", "quality", "adequate"],
    "big": ["large", "substantial", "significant"],
    "small": ["minimal", "limited", "minor"],
    "fast": ["quick", "rapid", "efficient"],
    "easy": ["simple", "straightforward", "uncomplicated"],
    "difficult": ["complex", "challenging", "hard"],
    "also": ["additionally", "furthermore", "moreover"],
    "because": ["since", "as", "due to"],
    "but": ["however", "nevertheless", "yet"],
    "with": ["using", "through", "by means of"],
    "for": ["to", "in order to", "intended for"],
    "from": ["originating", "derived", "starting"],
    "in": ["within", "inside", "during"],
    "or": ["alternatively", "otherwise", "or else"],
    "and": ["as well as", "along with", "plus"],
}

# ──────────────────────────────────────────────────────────────────────────────
# UTILITY
# ──────────────────────────────────────────────────────────────────────────────

def format_size(byte_count: int) -> str:
    if byte_count == 0:
        return "0 B"
    units = ["B", "KB", "MB", "GB"]
    i = 0
    val = float(byte_count)
    while val >= 1024 and i < len(units) - 1:
        val /= 1024
        i += 1
    return f"{val:.2f} {units[i]}"


def generate_session_id() -> str:
    return f"sess_{int(time.time() * 1000)}_{uuid.uuid4().hex[:9]}"


def ensure_dir(path: Path) -> Path:
    path.mkdir(parents=True, exist_ok=True)
    return path


# ──────────────────────────────────────────────────────────────────────────────
# SESSION MANAGEMENT
# ──────────────────────────────────────────────────────────────────────────────

def get_session_dir(session_id: str) -> Path:
    return UPLOADS_DIR / session_id


def init_session(session_id: str) -> dict:
    d = get_session_dir(session_id)
    ensure_dir(d / "original")
    ensure_dir(d / "processed")
    ensure_dir(d / "chunks")
    return {"session_id": session_id, "path": str(d)}


def list_sessions() -> list:
    sessions = []
    if not UPLOADS_DIR.exists():
        return sessions
    for entry in UPLOADS_DIR.iterdir():
        if not entry.is_dir():
            continue
        stat = entry.stat()
        size = _calc_dir_size(entry)
        file_count = sum(1 for f in entry.rglob("*") if f.is_file())
        meta_path = entry / "metadata.json"
        metadata = None
        if meta_path.exists():
            try:
                metadata = json.loads(meta_path.read_text("utf-8"))
            except Exception:
                pass
        sessions.append({
            "id": entry.name,
            "created_at": datetime.fromtimestamp(stat.st_ctime).strftime("%Y-%m-%d %H:%M"),
            "modified_at": datetime.fromtimestamp(stat.st_mtime).strftime("%Y-%m-%d %H:%M"),
            "modified_ts": stat.st_mtime,
            "file_size": size,
            "file_size_fmt": format_size(size),
            "file_count": file_count,
            "metadata": metadata,
        })
    sessions.sort(key=lambda x: x["modified_ts"], reverse=True)
    return sessions


def get_session_detail(session_id: str) -> Optional[dict]:
    session_path = get_session_dir(session_id)
    if not session_path.exists():
        return None
    structure = {}
    for subdir in ["original", "processed", "chunks"]:
        sub = session_path / subdir
        if sub.exists():
            structure[subdir] = [
                {
                    "name": f.name,
                    "size": f.stat().st_size,
                    "size_fmt": format_size(f.stat().st_size),
                    "modified": datetime.fromtimestamp(f.stat().st_mtime).strftime("%Y-%m-%d %H:%M"),
                }
                for f in sub.iterdir()
                if f.is_file()
            ]
    meta_path = session_path / "metadata.json"
    metadata = json.loads(meta_path.read_text("utf-8")) if meta_path.exists() else None
    return {"session_id": session_id, "metadata": metadata, "structure": structure}


def delete_session(session_id: str) -> dict:
    session_path = get_session_dir(session_id)
    if not session_path.exists():
        return {"success": False, "error": "Session not found"}
    size = _calc_dir_size(session_path)
    shutil.rmtree(session_path)
    return {"success": True, "freed_bytes": size, "freed_fmt": format_size(size)}


def cleanup_old_sessions(max_age_hours: int = 24, exclude_id: str = "") -> dict:
    max_age_ms = max_age_hours * 3600 * 1000
    now = time.time() * 1000
    deleted, freed = 0, 0
    if not UPLOADS_DIR.exists():
        return {"deleted": 0, "freed_bytes": 0, "freed_fmt": "0 B"}
    for entry in UPLOADS_DIR.iterdir():
        if not entry.is_dir():
            continue
        if entry.name == exclude_id:
            continue
        age_ms = now - entry.stat().st_mtime * 1000
        if age_ms > max_age_ms:
            size = _calc_dir_size(entry)
            shutil.rmtree(entry)
            freed += size
            deleted += 1
    return {"deleted": deleted, "freed_bytes": freed, "freed_fmt": format_size(freed)}


# ══════════════════════════════════════════════════════════════════════════════
# Hapus SEMUA session (kecuali session aktif)
# ══════════════════════════════════════════════════════════════════════════════

def delete_all_sessions(exclude_id: str = "") -> dict:
    """
    Delete ALL sessions from disk, except the one matching exclude_id.
    Returns count of deleted sessions and total freed bytes.
    """
    deleted, freed = 0, 0
    if not UPLOADS_DIR.exists():
        return {"deleted": 0, "freed_bytes": 0, "freed_fmt": "0 B"}
    for entry in list(UPLOADS_DIR.iterdir()):
        if not entry.is_dir():
            continue
        if entry.name == exclude_id:
            continue
        size = _calc_dir_size(entry)
        shutil.rmtree(entry)
        freed += size
        deleted += 1
    return {"deleted": deleted, "freed_bytes": freed, "freed_fmt": format_size(freed)}


def _calc_dir_size(path: Path) -> int:
    return sum(f.stat().st_size for f in path.rglob("*") if f.is_file())


# ──────────────────────────────────────────────────────────────────────────────
# FILE SAVING
# ──────────────────────────────────────────────────────────────────────────────

def save_uploaded_file(session_id: str, uploaded_file) -> dict:
    """Save a Streamlit UploadedFile to disk. Returns file info dict."""
    session_dir = get_session_dir(session_id)
    original_dir = ensure_dir(session_dir / "original")

    safe_name = re.sub(r"[^a-zA-Z0-9._\-]", "_", uploaded_file.name)
    unique_name = f"{int(time.time() * 1000)}_{safe_name}"
    dest = original_dir / unique_name

    content = uploaded_file.read()
    dest.write_bytes(content)

    return {
        "original_name": uploaded_file.name,
        "saved_name": unique_name,
        "path": str(dest),
        "size": len(content),
        "size_fmt": format_size(len(content)),
        "ext": Path(uploaded_file.name).suffix.lower(),
        "session_id": session_id,
    }


# ──────────────────────────────────────────────────────────────────────────────
# TEXT EXTRACTION  (PDF, DOC, DOCX, TXT, MD, Images)
# ──────────────────────────────────────────────────────────────────────────────

def extract_text_from_file(file_path: str, file_ext: str) -> str:
    """Extract raw text from PDF / DOC / DOCX / TXT / MD / image files."""
    p = Path(file_path)
    ext = file_ext.lower()

    # ── Plain text ──
    if ext in (".txt", ".md"):
        for enc in ("utf-8", "latin-1", "cp1252"):
            try:
                return p.read_text(enc)
            except UnicodeDecodeError:
                continue
        return p.read_bytes().decode("utf-8", errors="replace")

    # ── PDF ──
    elif ext == ".pdf":
        try:
            import pypdf
            reader = pypdf.PdfReader(str(p))
            pages = []
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    pages.append(text)
            extracted = "\n\n".join(pages)
            if len(extracted.strip()) >= 50:
                return extracted
        except ImportError:
            pass
        return _extract_pdf_fallback(p)

    # ── DOCX ──
    elif ext == ".docx":
        try:
            import docx as python_docx
            doc = python_docx.Document(str(p))
            parts = [para.text for para in doc.paragraphs if para.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            parts.append(cell.text.strip())
            return "\n\n".join(parts)
        except ImportError:
            return _extract_docx_fallback(p)

    # ── DOC (legacy) ──
    elif ext == ".doc":
        return _extract_doc_legacy(p)

    # ── Excel XLSX ──
    elif ext == ".xlsx":
        return _extract_excel_text(p, engine="openpyxl")

    # ── Excel XLS (legacy) ──
    elif ext == ".xls":
        return _extract_excel_text(p, engine="xlrd")

    # ── Images (OCR) ──
    elif ext in IMAGE_EXTENSIONS:
        return _extract_image_text(p)

    return ""


def _extract_pdf_fallback(p: Path) -> str:
    try:
        import subprocess
        result = subprocess.run(
            ["pdftotext", str(p), "-"],
            capture_output=True, text=True, timeout=120
        )
        if result.returncode == 0 and result.stdout.strip():
            return result.stdout
    except Exception:
        pass
    return f"[Cannot extract PDF: {p.name}. Install pypdf or pdftotext]"


def _extract_docx_fallback(p: Path) -> str:
    try:
        import zipfile
        import xml.etree.ElementTree as ET
        texts = []
        with zipfile.ZipFile(str(p)) as z:
            with z.open("word/document.xml") as f:
                tree = ET.parse(f)
                ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
                for t in tree.findall(".//w:t", ns):
                    if t.text:
                        texts.append(t.text)
        return " ".join(texts)
    except Exception:
        return f"[Cannot extract DOCX: {p.name}. Install python-docx]"


def _extract_doc_legacy(p: Path) -> str:
    """Extract text from legacy .doc files."""
    try:
        import subprocess
        result = subprocess.run(
            ["antiword", str(p)],
            capture_output=True, text=True, timeout=60
        )
        if result.returncode == 0 and result.stdout.strip():
            return result.stdout
    except Exception:
        pass

    try:
        import subprocess, tempfile
        with tempfile.TemporaryDirectory() as tmpdir:
            subprocess.run(
                ["libreoffice", "--headless", "--convert-to", "txt:Text", "--outdir", tmpdir, str(p)],
                capture_output=True, timeout=120
            )
            txt_file = Path(tmpdir) / (p.stem + ".txt")
            if txt_file.exists():
                return txt_file.read_text("utf-8", errors="replace")
    except Exception:
        pass

    try:
        raw = p.read_bytes()
        texts = re.findall(rb"[\x20-\x7e]{4,}", raw)
        return " ".join(t.decode("ascii", errors="replace") for t in texts)
    except Exception:
        return f"[Cannot extract DOC: {p.name}. Install antiword or LibreOffice]"


def _extract_excel_text(p: Path, engine: str = "openpyxl") -> str:
    """Extract text from Excel files (.xlsx with openpyxl, .xls with xlrd)."""
    try:
        import pandas as pd
        xl = pd.ExcelFile(str(p), engine=engine)
        parts = []
        for sheet_name in xl.sheet_names:
            df = xl.parse(sheet_name)
            if df.empty:
                continue
            parts.append(f"=== Sheet: {sheet_name} ===")
            # Header row
            parts.append("\t".join(str(c) for c in df.columns))
            # Data rows
            for _, row in df.iterrows():
                row_text = "\t".join("" if (str(v) == "nan") else str(v) for v in row)
                if row_text.strip():
                    parts.append(row_text)
        return "\n".join(parts) if parts else f"[Empty Excel file: {p.name}]"
    except ImportError as e:
        missing = "openpyxl" if engine == "openpyxl" else "xlrd"
        return f"[Cannot extract Excel: install {missing}. Error: {e}]"
    except Exception as e:
        return f"[Cannot extract Excel: {p.name}. Error: {e}]"


def _extract_image_text(p: Path) -> str:
    """Extract text from image using pytesseract OCR."""
    try:
        import pytesseract
        from PIL import Image
        img = Image.open(str(p))
        if img.mode not in ("RGB", "L"):
            img = img.convert("RGB")
        try:
            text = pytesseract.image_to_string(img, lang="ind+eng")
        except Exception:
            text = pytesseract.image_to_string(img)
        return text.strip()
    except Exception as e:
        return f"[Image: {p.name} — Install pytesseract + Tesseract OCR for text extraction. Error: {e}]"


# ──────────────────────────────────────────────────────────────────────────────
# CHUNKING
# ──────────────────────────────────────────────────────────────────────────────

def clean_text(text: str, remove_extra_spaces: bool = True, remove_special: bool = False) -> str:
    if remove_extra_spaces:
        text = re.sub(r"\s+", " ", text).strip()
        text = re.sub(r"\n{3,}", "\n\n", text)
    if remove_special:
        text = re.sub(r"[^\w\s.,!?;:()\-\'\"\n]", "", text)
    return text.strip()


def chunk_text(
    text: str,
    chunk_size: int = 512,
    overlap: int = 50,
    method: str = "tokens",
) -> list:
    text = text.strip()
    if not text:
        return []

    if method == "paragraphs":
        paras = [p.strip() for p in re.split(r"\n\s*\n", text) if p.strip()]
        chunks, current = [], ""
        for para in paras:
            if len((current + " " + para).split()) <= chunk_size:
                current = (current + " " + para).strip()
            else:
                if current:
                    chunks.append(current)
                current = para
        if current:
            chunks.append(current)
        return chunks

    elif method == "sentences":
        sentences = re.split(r"(?<=[.!?])\s+", text)
        chunks, current_words = [], []
        for sent in sentences:
            words = sent.split()
            if len(current_words) + len(words) <= chunk_size:
                current_words.extend(words)
            else:
                if current_words:
                    chunks.append(" ".join(current_words))
                current_words = current_words[-overlap:] + words if overlap else words
        if current_words:
            chunks.append(" ".join(current_words))
        return chunks

    else:  # tokens (word-based)
        words = text.split()
        chunks = []
        i = 0
        while i < len(words):
            end = min(i + chunk_size, len(words))
            chunks.append(" ".join(words[i:end]))
            i += chunk_size - overlap
            if i < 0:
                i = chunk_size
        return chunks


# ──────────────────────────────────────────────────────────────────────────────
# SNI DOCUMENT FILTER
# ──────────────────────────────────────────────────────────────────────────────

# ---------------------------------------------------------------------------
# 1. SECTION KEYWORDS — headings we WANT to keep
# ---------------------------------------------------------------------------
_SNI_SECTION_KEYWORDS = [
    r"ruang\s+lingkup", r"\bscope\b",
    r"persyaratan", r"\bmutu\b", r"\brequirement[s]?\b", r"\bquality\b",
    r"syarat\s+mutu", r"spesifikasi", r"\bspecification[s]?\b",
    r"pengujian", r"metode\s+uji", r"\buji\b", r"\btesting\b",
    r"test\s+method[s]?", r"\bsampling\b",
    r"verifikasi", r"\bverification\b", r"\bvalidation\b", r"validasi",
    r"istilah\s+dan\s+definisi", r"terms?\s+and\s+definitions?",
    r"acuan\s+normatif", r"normative\s+references?",
    r"simbol", r"lambang", r"singkatan", r"\bsymbol[s]?\b",
    r"cara\s+pengambilan", r"pengambilan\s+contoh",
    r"penandaan", r"pelabelan", r"\bmarking\b", r"\blabelling?\b",
    r"higiene", r"\bhygiene\b", r"keselamatan", r"\bsafety\b",
    r"pengemasan", r"\bpackaging\b",
    r"klasifikasi", r"\bclassification\b",
]

# ---------------------------------------------------------------------------
# 2. SECTION KEYWORDS — headings (and their content) to EXCLUDE
# ---------------------------------------------------------------------------
_SNI_EXCLUDE_KEYWORDS = [
    # Daftar Isi / Table of Contents
    r"daftar\s+isi", r"table\s+of\s+contents?", r"^contents?$",

    # Prakata / Foreword / Kata Pengantar
    r"\bprakata\b", r"\bforeword\b", r"kata\s+pengantar", r"kata\s+sambutan",
    r"sambutan\s+(?:ketua|kepala|direktur)", r"\bpreface\b",

    # Pendahuluan / Introduction
    r"\bpendahuluan\b", r"\bintroduction\b",

    # Bibliografi / Daftar Pustaka / References
    r"\bbibliografi\b", r"\bbibliography\b",
    r"daftar\s+pustaka", r"\breferensi\b", r"\breferences?\b",
    r"daftar\s+acuan", r"sumber\s+pustaka",

    # Informasi Perumus / Komite Teknis
    r"informasi\s+perumus", r"komite\s+teknis", r"drafting\s+committee",
    r"panitia\s+teknis", r"technical\s+committee",
    r"perumus(?:an)?\s+standar", r"penyusun(?:an)?\s+standar",
    r"keanggotaan\s+komite", r"committee\s+membership",
    r"susunan\s+keanggotaan", r"daftar\s+anggota\s+komite",

    # Lampiran Informatif / Informative Annex
    r"lampiran\s+(?:informasi|informative[a-z]*)",
    r"annex\s+(?:informative[a-z]*|informasi)",

    # Copyright / Hak Cipta
    r"hak\s+cipta\b", r"\bcopyright\b", r"ketentuan\s+penggunaan",
]

# ---------------------------------------------------------------------------
# 3. NORMATIVE SENTENCE MARKERS — always kept regardless of section
# ---------------------------------------------------------------------------
_NORMATIVE_PATTERNS = [
    r"\bharus\b", r"\bshall\b",
    r"\bsebaiknya\b", r"\bshould\b",
]

# ---------------------------------------------------------------------------
# 4. NOISE LINES — stripped from ALL sections (footer, copyright, page nos)
# ---------------------------------------------------------------------------
_NOISE_LINE_RE = re.compile(
    r"(?:"
    # copyright symbols and phrases
    r"©"
    r"|hak\s+cipta"
    r"|all\s+rights?\s+reserved"
    r"|dilarang\s+memperbanyak"
    r"|dilarang\s+mendistribusikan"
    r"|dilarang\s+mereproduksi"
    r"|tanpa\s+izin\s+tertulis"
    r"|reproduction\s+prohibited"
    r"|hak\s+penerbitan"
    # BSN institutional lines
    r"|badan\s+standardisasi\s+nasional"
    r"|standardisasi\s+nasional\s+indonesia"
    r"|dokinfo@bsn"
    r"|www\.bsn\.go\.id"
    r"|diterbitkan\s+di\s+jakarta"
    r"|diterbitkan\s+oleh"
    r"|email\s*:"
    r"|telp\s*[\.:]"
    r"|fax\s*[\.:]"
    # ICS classification lines (e.g. "ICS 67.220.10")
    r"|^\s*ICS\s+\d"
    # Lone page numbers
    r"|^\s*\d{1,4}\s*$"
    # Page labels
    r"|halaman\s+\d"
    r"|page\s+\d+\s+of\s+\d+"
    r")",
    re.IGNORECASE,
)

# ---------------------------------------------------------------------------
# 5. COVER / PRE-CONTENT NOISE — lines seen on cover and copyright pages
#    that should never appear in output even if found in "title" pass
# ---------------------------------------------------------------------------
_COVER_NOISE_RE = re.compile(
    r"(?:"
    r"standar\s+nasional\s+indonesia"       # cover banner
    r"|national\s+standard\s+of\s+indonesia"
    r"|^\s*ICS\b"                            # ICS line on cover
    r"|^\s*SNI\s+\d[\w\.\-:]*\s*$"          # bare SNI number line (cover stamp)
    r"|badan\s+standardisasi"
    r"|diterbitkan"
    r"|tanpa\s+izin"
    r"|dilarang"
    r"|bentuk\s+apapun"
    r"|secara\s+elektronik"
    r"|pangkal\s+bsn"
    r"|tercetak"
    r"|dokinfo"
    r"|www\.bsn"
    r"|email\s*:"
    r"|hak\s+cipta"
    r"|©"
    r"|all\s+rights?\s+reserved"
    r"|reproduction\s+prohibited"
    r")",
    re.IGNORECASE,
)

# ---------------------------------------------------------------------------
# 6. COMPILED REGEXES
# ---------------------------------------------------------------------------
_SNI_SECTION_RE  = re.compile("|".join(_SNI_SECTION_KEYWORDS),  re.IGNORECASE)
_SNI_EXCLUDE_RE  = re.compile("|".join(_SNI_EXCLUDE_KEYWORDS),  re.IGNORECASE)
_NORMATIVE_RE    = re.compile("|".join(_NORMATIVE_PATTERNS),    re.IGNORECASE)

# SNI ID pattern — strict: "SNI 4:2025", "SNI 01-3140-2010", "SNI ISO 9001:2015"
# Must be at start of line or after whitespace, and look like a primary identifier
# (not a mid-sentence citation reference)
_SNI_ID_LINE_RE = re.compile(
    r"^\s*SNI(?:\s+(?:ISO|IEC|ASTM|EN|BS))?[\s\-]+\d[\w\.\-:/]*(?:[:\-]\d[\w\.\-:/]*)?\s*$",
    re.IGNORECASE,
)
# For page-header extraction: short standalone SNI line (≤ 30 chars)
_SNI_HEADER_RE = re.compile(
    r"^SNI(?:\s+(?:ISO|IEC|ASTM|EN|BS))?[\s\-]+[\w\.\-:/]+(?:[\s\-:]\d[\w\.\-:/]*)?\s*$",
    re.IGNORECASE,
)

# Heading: numbered section "1.", "2.1.", up to 55 chars after the number
_HEADING_RE = re.compile(r"^\d{1,2}(?:\.\d+)*\.?\s+\S.{0,54}$")

# Unnumbered heading: short line (2–70 chars) — used for keyword-matched headings
_UNNUMBERED_KW_RE = re.compile(r"^[\w][^\n]{1,68}$")

# ToC entry pattern: line with trailing dots and/or page number
_TOC_LINE_RE = re.compile(r"\.{3,}|\s{3,}\d+\s*$")


# ---------------------------------------------------------------------------
# HELPERS
# ---------------------------------------------------------------------------

def _is_noise(line: str) -> bool:
    return bool(_NOISE_LINE_RE.search(line))


def _is_excluded_heading(s: str) -> bool:
    return bool(_SNI_EXCLUDE_RE.search(s))


def _is_relevant_heading(s: str) -> bool:
    return bool(_SNI_SECTION_RE.search(s))


def _is_numbered_heading(s: str) -> bool:
    return bool(_HEADING_RE.match(s))


def _is_section_heading(s: str) -> bool:
    """True for any line that acts as a section boundary (numbered or keyword-matched)."""
    if _HEADING_RE.match(s):
        return True
    # Any short line that matches exclude OR keep keyword list is a section boundary
    if len(s) <= 70 and (
        _SNI_EXCLUDE_RE.search(s) or _SNI_SECTION_RE.search(s)
    ):
        return True
    return False


def _extract_sni_id_from_headers(lines: list[str]) -> str | None:
    """
    Scan repeating page-header lines for the primary SNI identifier.
    In BSN PDFs the SNI number (e.g. "SNI 4:2025") often appears as a
    running header on every content page.  We pick the most-frequent
    short SNI-like line as the canonical ID.
    """
    from collections import Counter
    candidates = []
    for ln in lines:
        s = ln.strip()
        if _SNI_HEADER_RE.match(s) and len(s) <= 30:
            candidates.append(s)
    if not candidates:
        return None
    # Most common short SNI header line = primary identifier
    most_common, count = Counter(candidates).most_common(1)[0]
    # Only trust it if it appears more than once (true running header)
    return most_common if count > 1 else (most_common if candidates else None)


def _extract_document_title(lines: list[str], sni_id: str | None) -> str | None:
    """
    Extract the document title (commodity/subject name) from the content area.
    Strategy: look for the first non-boilerplate, non-SNI-id short line that
    appears right after a page with the SNI id header, before section 1.
    """
    sni_id_norm = (sni_id or "").strip().lower()
    # Collect candidate title lines — short, clean, not noise, not the SNI id itself
    for ln in lines:
        s = ln.strip()
        if not s:
            continue
        if _is_noise(s) or _COVER_NOISE_RE.search(s):
            continue
        if _is_section_heading(s):
            break
        # Skip the SNI number line itself
        if sni_id_norm and s.lower() == sni_id_norm:
            continue
        if _SNI_ID_LINE_RE.match(s):
            continue
        # Skip very long lines (copyright paragraphs etc.)
        if len(s) > 120:
            continue
        # A plausible title: 3–80 chars, mostly letters
        letter_ratio = sum(c.isalpha() for c in s) / max(len(s), 1)
        if 3 <= len(s) <= 80 and letter_ratio >= 0.5:
            return s
    return None


# ---------------------------------------------------------------------------
# MAIN FILTER
# ---------------------------------------------------------------------------

def extract_sni_sections(raw_text: str) -> str:
    """
    Filter raw text extracted from an SNI PDF.

    OUTPUT structure (plain text, no === markers):
        <no_sni>: <judul>

        <section heading>
        <content lines>
        ...

    EXCLUDED (completely removed):
      • Cover page (everything before section 1, except title extraction)
      • Copyright / ICS / BSN institutional lines (any page)
      • Prakata / Foreword / Kata Pengantar / Kata Sambutan
      • Pendahuluan / Introduction
      • Daftar Isi / Table of Contents
      • Bibliografi / Daftar Pustaka / References
      • Informasi Perumus / Komite Teknis / Panitia Teknis
      • Footer noise (page numbers, BSN address, email, URL)

    RETAINED:
      • no_sni + judul header line
      • Ruang Lingkup / Scope
      • Acuan Normatif / Normative References
      • Istilah dan Definisi / Terms and Definitions
      • Persyaratan / Mutu / Spesifikasi / Requirement
      • Pengujian / Metode Uji / Verifikasi / Validasi / Sampling
      • Klasifikasi / Penandaan / Pengemasan / Keselamatan / Higiene
      • Any sentence containing: harus / shall / sebaiknya / should
        (even when inside an otherwise-excluded section)
    """
    lines = raw_text.splitlines()

    # ── Step 1: extract SNI ID from running page headers ────────────────────
    sni_id = _extract_sni_id_from_headers(lines)

    # ── Step 2: find where content sections start (first numbered heading) ──
    first_section_idx = None
    for i, line in enumerate(lines):
        s = line.strip()
        if s and _is_numbered_heading(s):
            first_section_idx = i
            break

    # ── Step 3: extract document title from lines before section 1 ──────────
    title_candidates = lines[:first_section_idx] if first_section_idx else lines[:60]
    doc_title = _extract_document_title(title_candidates, sni_id)

    # ── Step 4: walk sections (start from section 1 or beginning) ───────────
    kept_sections: list[tuple[str, list[str]]] = []
    current_label: str | None = None
    current_lines: list[str] = []
    in_relevant = False
    in_excluded = False

    start = first_section_idx if first_section_idx is not None else 0

    for line in lines[start:]:
        s = line.strip()

        # Always drop noise and cover-noise lines
        if not s or _is_noise(s) or _COVER_NOISE_RE.search(s):
            if in_relevant and current_lines:
                current_lines.append("")
            continue

        # Drop ToC-style entries (dotted leaders or trailing page numbers)
        if _TOC_LINE_RE.search(s):
            continue

        is_heading = _is_section_heading(s)

        if is_heading:
            # Flush previous section
            if current_label is not None and in_relevant and current_lines:
                kept_sections.append((current_label, _clean_section_lines(current_lines)))

            current_label = s
            current_lines = []

            if _is_excluded_heading(s):
                in_excluded = True
                in_relevant = False
            elif _is_relevant_heading(s):
                in_excluded = False
                in_relevant = True
            else:
                in_excluded = False
                in_relevant = False

        else:
            if in_relevant:
                # Skip lines that are just the SNI id repeated (running header)
                if sni_id and s.lower() == sni_id.lower():
                    continue
                current_lines.append(s)
            elif not in_excluded and _NORMATIVE_RE.search(s):
                # Keep normative sentences only outside of excluded sections
                kept_sections.append(("[normative]", [s]))

    # Flush last section
    if current_label is not None and in_relevant and current_lines:
        kept_sections.append((current_label, _clean_section_lines(current_lines)))

    # ── Step 5: assemble output ──────────────────────────────────────────────
    parts: list[str] = []

    # Header line: "no_sni: judul"
    header_parts = []
    if sni_id:
        header_parts.append(sni_id)
    if doc_title:
        header_parts.append(doc_title)
    if header_parts:
        parts.append(": ".join(header_parts))
        parts.append("")

    for label, content_lines in kept_sections:
        if label != "[normative]":
            parts.append(label)
        parts.extend(content_lines)
        parts.append("")

    return "\n".join(parts).strip()


def _clean_section_lines(lines: list[str]) -> list[str]:
    """Remove trailing blank lines and deduplicate adjacent blanks."""
    out = []
    prev_blank = False
    for ln in lines:
        if ln == "":
            if not prev_blank:
                out.append("")
            prev_blank = True
        else:
            out.append(ln)
            prev_blank = False
    # strip leading/trailing blanks
    while out and out[0] == "":
        out.pop(0)
    while out and out[-1] == "":
        out.pop()
    return out


def _process_single_file(args) -> dict:
    """Process a single file — used by parallel executor."""
    info, settings = args
    chunk_size = int(settings.get("chunk_size", 512))
    overlap = int(settings.get("overlap", 50))
    method = settings.get("chunk_method", "tokens")
    remove_spaces = settings.get("remove_extra_spaces", True)
    remove_special = settings.get("remove_special_chars", False)
    source_tag = settings.get("source_tag", "")
    output_format = settings.get("output_format", "training")

    try:
        raw = extract_text_from_file(info["path"], info["ext"])
        if not raw.strip():
            return {
                "file_result": {**info, "status": "empty", "chunks": 0},
                "items": [],
                "chars": 0,
            }

        # Apply SNI section filter for PDF files
        if info["ext"].lower() == ".pdf":
            raw = extract_sni_sections(raw)
            if not raw.strip():
                return {
                    "file_result": {**info, "status": "empty", "chunks": 0,
                                    "note": "No relevant SNI sections found"},
                    "items": [],
                    "chars": 0,
                }

        cleaned = clean_text(raw, remove_spaces, remove_special)
        chunks = chunk_text(cleaned, chunk_size, overlap, method)
        items = []

        for idx, chunk in enumerate(chunks):
            if output_format == "qa":
                item = {"instruction": "", "input": chunk, "output": ""}
            elif output_format == "messages":
                item = {
                    "messages": [
                        {"role": "system", "content": "You are a helpful assistant."},
                        {"role": "user", "content": chunk},
                        {"role": "assistant", "content": ""},
                    ]
                }
            else:
                item = {
                    "id": f"{info['session_id']}_tmp_{idx:04d}",
                    "text": chunk,
                    "source": source_tag or info["original_name"],
                    "chunk_index": idx + 1,
                    "total_chunks_in_doc": len(chunks),
                }
            items.append(item)

        return {
            "file_result": {**info, "status": "success", "chunks": len(chunks)},
            "items": items,
            "chars": len(cleaned),
        }

    except Exception as e:
        return {
            "file_result": {**info, "status": "error", "error": str(e), "chunks": 0},
            "items": [],
            "chars": 0,
        }


def process_documents(
    file_infos: list,
    settings: dict,
    max_workers: int = 12,
    progress_callback=None,
) -> dict:
    """Process a list of saved file infos into training data using parallel workers."""
    all_results = [None] * len(file_infos)
    total = len(file_infos)
    completed_count = [0]

    with ThreadPoolExecutor(max_workers=max_workers) as executor:
        future_to_idx = {
            executor.submit(_process_single_file, (info, settings)): i
            for i, info in enumerate(file_infos)
        }
        for future in as_completed(future_to_idx):
            idx = future_to_idx[future]
            all_results[idx] = future.result()
            completed_count[0] += 1
            if progress_callback:
                progress_callback("extract", completed_count[0], total, file_infos[idx]["original_name"])

    all_items = []
    file_results = []
    stats = {
        "total_files": len(file_infos),
        "processed_files": 0,
        "failed_files": 0,
        "total_chunks": 0,
        "total_chars": 0,
    }

    for res in all_results:
        if res is None:
            continue
        file_results.append(res["file_result"])
        all_items.extend(res["items"])
        if res["file_result"]["status"] == "success":
            stats["processed_files"] += 1
            stats["total_chunks"] += res["file_result"]["chunks"]
            stats["total_chars"] += res["chars"]
        else:
            stats["failed_files"] += 1

    session_id = file_infos[0]["session_id"] if file_infos else "sess"
    for i, item in enumerate(all_items):
        if isinstance(item, dict) and "id" in item:
            item["id"] = f"{session_id}_{i + 1:04d}"

    return {
        "version": "1.0",
        "type": "training",
        "created_at": datetime.now(timezone.utc).isoformat(),
        "settings": settings,
        "stats": stats,
        "file_results": file_results,
        "total": len(all_items),
        "total_chunks": len(all_items),
        "data": all_items,
    }


# ──────────────────────────────────────────────────────────────────────────────
# PARAPHRASE ENGINE (WITHOUT API)
# ──────────────────────────────────────────────────────────────────────────────

# Words that must NEVER be replaced during paraphrasing (normative/modal terms)
PROTECTED_WORDS = {
    # English normative
    "shall", "should", "must", "may", "can", "will", "would",
    # Indonesian normative
    "harus", "sebaiknya", "wajib", "boleh", "dapat", "tidak boleh",
    "dilarang", "dipersyaratkan", "direkomendasikan", "diperbolehkan",
}


def _replace_synonyms(text: str, synonyms_dict: dict, replace_prob: float = 0.7) -> str:
    """Replace words with their synonyms based on probability, preserving protected words."""
    words = text.split()
    for i, word in enumerate(words):
        clean_word = re.sub(r'[^\w]', '', word.lower())
        # Skip protected normative/modal words — never replace them
        if clean_word in PROTECTED_WORDS:
            continue
        if clean_word in synonyms_dict and random.random() < replace_prob:
            synonyms = synonyms_dict[clean_word]
            synonym = random.choice(synonyms)
            punctuation = re.findall(r'[^\w]', word)
            if punctuation:
                words[i] = synonym + punctuation[-1] if punctuation[-1] else synonym
            else:
                words[i] = synonym
    return " ".join(words)


def _rephrase_sentence_structure(text: str) -> str:
    """Slightly modify sentence structure without changing meaning."""
    sentences = re.split(r'(?<=[.!?])\s+', text)
    rephrased = []

    for sentence in sentences:
        if len(sentence.split()) < 5:
            rephrased.append(sentence)
            continue

        # Don't restructure sentences containing normative/modal words
        sentence_lower = sentence.lower()
        has_protected = any(pw in sentence_lower.split() or
                            f" {pw} " in f" {sentence_lower} "
                            for pw in PROTECTED_WORDS)
        if has_protected:
            rephrased.append(sentence)
            continue

        if ',' in sentence:
            parts = sentence.split(',', 1)
            if len(parts) == 2 and random.random() < 0.3:
                sentence = f"{parts[1].strip()} {parts[0].strip()}"
                sentence = sentence[0].upper() + sentence[1:]

        rephrased.append(sentence)

    return " ".join(rephrased)


def _add_transitional_phrases(text: str) -> str:
    """Add or modify transitional phrases."""
    indonesian_transitions = [
        "Selain itu, ", "Sebagai tambahan, ", "Di sisi lain, ",
        "Dalam konteks ini, ", "Sehubungan dengan itu, ", "Pada dasarnya, "
    ]
    english_transitions = [
        "Additionally, ", "Furthermore, ", "On the other hand, ",
        "In this context, ", "Related to this, ", "Essentially, "
    ]

    if len(text.split()) < 30:
        return text

    if random.random() < 0.2:
        if any(word in text.lower() for word in ["dan", "yang", "dengan", "untuk", "dari"]):
            transition = random.choice(indonesian_transitions)
        else:
            transition = random.choice(english_transitions)
        text = transition + text

    return text


def paraphrase_text(text: str, language: str = "auto") -> str:
    """
    Paraphrase text without using an API.
    Default language is 'auto' for auto-detection.
    """
    if not text.strip():
        return text

    if language == "id":
        synonyms = INDONESIAN_SYNONYMS
    elif language == "en":
        synonyms = ENGLISH_SYNONYMS
    else:
        # Auto-detect language (simple heuristic)
        if any(word in text.lower() for word in ["dan", "yang", "dengan", "untuk", "dari"]):
            synonyms = INDONESIAN_SYNONYMS
        else:
            synonyms = ENGLISH_SYNONYMS

    paraphrased = _replace_synonyms(text, synonyms)
    paraphrased = _rephrase_sentence_structure(paraphrased)
    paraphrased = _add_transitional_phrases(paraphrased)

    return paraphrased


def _paraphrase_item_task(args) -> tuple:
    """Returns (index, paraphrased_item)."""
    import copy
    idx, item, language = args
    item = copy.deepcopy(item)
    try:
        if "text" in item:
            item["text"] = paraphrase_text(item["text"], language)
        elif "input" in item:
            item["input"] = paraphrase_text(item["input"], language)
        elif "messages" in item:
            for msg in item["messages"]:
                if msg.get("role") == "user" and msg.get("content"):
                    msg["content"] = paraphrase_text(msg["content"], language)
    except Exception:
        pass
    return idx, item


def paraphrase_dataset(
    result: dict,
    language: str = "auto",
    progress_callback=None,
    max_workers: int = 10,
) -> dict:
    """
    Paraphrase every text field in result['data'] using parallel processing.
    Default language is 'auto' for auto-detection.
    """
    import copy
    new_result = copy.deepcopy(result)
    items = new_result.get("data", [])
    total = len(items)
    paraphrased_items = [None] * total

    completed_count = [0]
    tasks = [(i, item, language) for i, item in enumerate(items)]

    with ThreadPoolExecutor(max_workers=max_workers) as executor:
        futures = {executor.submit(_paraphrase_item_task, task): task[0] for task in tasks}
        for future in as_completed(futures):
            idx, para_item = future.result()
            paraphrased_items[idx] = para_item
            completed_count[0] += 1
            if progress_callback:
                progress_callback(completed_count[0], total)

    if progress_callback:
        progress_callback(total, total)

    new_result["data"] = paraphrased_items
    new_result["paraphrased"] = True
    new_result["paraphrase_language"] = language
    new_result["paraphrased_at"] = datetime.now(timezone.utc).isoformat()
    return new_result


def process_and_paraphrase(
    file_infos: list,
    settings: dict,
    language: str = "auto",
    progress_callback=None,
) -> dict:
    """
    One-shot: extract text from files, chunk, then immediately paraphrase.
    Default language is 'auto' for auto-detection.
    progress_callback(stage: str, current: int, total: int)
    """
    if progress_callback:
        progress_callback("extract", 0, len(file_infos), "")

    def _extract_cb(stage, cur, tot, fname=""):
        if progress_callback:
            progress_callback(stage, cur, tot, fname)

    result = process_documents(file_infos, settings, max_workers=12, progress_callback=_extract_cb)

    if progress_callback:
        progress_callback("extract", len(file_infos), len(file_infos), "")

    def _para_cb(current, total):
        if progress_callback:
            progress_callback("paraphrase", current, total, "")

    para_result = paraphrase_dataset(result, language, _para_cb, max_workers=10)
    return para_result


# ──────────────────────────────────────────────────────────────────────────────
# SAVE PROCESSED DATA
# ──────────────────────────────────────────────────────────────────────────────

def save_processed_data(session_id: str, data: dict, settings: dict) -> dict:
    session_dir = get_session_dir(session_id)
    processed_dir = ensure_dir(session_dir / "processed")
    chunks_dir = ensure_dir(session_dir / "chunks")

    output_path = processed_dir / "output.json"
    output_path.write_text(json.dumps(data, ensure_ascii=False, indent=2), "utf-8")

    chunk_count = 0
    if isinstance(data.get("data"), list):
        def _write_chunk(args):
            i, item = args
            (chunks_dir / f"chunk_{i + 1:03d}.json").write_text(
                json.dumps(item, ensure_ascii=False, indent=2), "utf-8"
            )
            return i

        with ThreadPoolExecutor(max_workers=20) as executor:
            results = list(executor.map(_write_chunk, enumerate(data["data"])))
            chunk_count = len(results)

    metadata = {
        "session_id": session_id,
        "saved_at": datetime.now(timezone.utc).isoformat(),
        "settings": settings,
        "stats": {
            "total_chunks": chunk_count,
            "output_file": str(output_path),
            "chunks_directory": str(chunks_dir),
        },
    }
    (session_dir / "metadata.json").write_text(
        json.dumps(metadata, ensure_ascii=False, indent=2), "utf-8"
    )

    return {
        "success": True,
        "processed_path": str(output_path),
        "chunks_saved": chunk_count,
    }


# ──────────────────────────────────────────────────────────────────────────────
# JSON MERGER
# ──────────────────────────────────────────────────────────────────────────────

def detect_merge_strategy(parsed_files: list, strategy: str = "auto") -> str:
    if strategy != "auto":
        return strategy
    if all(isinstance(f["content"], dict) and isinstance(f["content"].get("data"), list)
           for f in parsed_files):
        return "data"
    if all(isinstance(f["content"], list) for f in parsed_files):
        return "array"
    return "object"


def merge_json_files(file_contents: list, strategy: str = "auto") -> dict:
    merge_mode = detect_merge_strategy(file_contents, strategy)
    stats = {
        "files_count": len(file_contents),
        "merge_mode": merge_mode,
        "item_counts": [],
    }

    if merge_mode == "data":
        all_items = []
        for f in file_contents:
            items = f["content"].get("data", [])
            all_items.extend(items)
            stats["item_counts"].append({"file": f["name"], "items": len(items)})
        merged = {
            "version": "1.0",
            "type": "training",
            "merged": True,
            "merged_at": datetime.now(timezone.utc).isoformat(),
            "source_files": [f["name"] for f in file_contents],
            "total": len(all_items),
            "total_chunks": len(all_items),
            "data": all_items,
        }

    elif merge_mode == "array":
        all_items = []
        for f in file_contents:
            c = f["content"]
            if isinstance(c, list):
                arr = c
            elif isinstance(c, dict) and isinstance(c.get("data"), list):
                arr = c["data"]
            elif isinstance(c, dict):
                arr = [c]
            else:
                arr = []
            all_items.extend(arr)
            stats["item_counts"].append({"file": f["name"], "items": len(arr)})
        merged = all_items

    else:
        merged = {
            "_merged": True,
            "_merged_at": datetime.now(timezone.utc).isoformat(),
            "_source_files": [f["name"] for f in file_contents],
        }
        for f in file_contents:
            src = f["content"] if isinstance(f["content"], dict) else {"_array_data": f["content"]}
            merged.update(src)
            stats["item_counts"].append({"file": f["name"], "keys": len(src)})

    return {"merged": merged, "merge_mode": merge_mode, "stats": stats}
